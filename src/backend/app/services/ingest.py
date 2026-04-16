"""
Ingestion Service - Process documents through the knowledge layer.

Pipeline:
1. Extract text from documents (PDFs, DOCX, TXT, etc.)
2. Chunk documents
3. Store chunks in SQLite
4. Embed and index in Qdrant (embedded, in-process)
5. Extract entities and create graph nodes/edges

Production Desktop Sidecar: SQLite + embedded Qdrant + bundled LLMs.
"""
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import uuid
import logging
import time
import asyncio
from concurrent.futures import ThreadPoolExecutor

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
from pypdf import PdfReader

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from sqlalchemy.orm import Session

from app.core.database import get_session, Node, Edge, Document, DocumentChunk
from app.core.config import settings
from app.services.llm import LLMService
from app.core.qdrant_store import get_qdrant_client
from app.services.docling_parser import DoclingParser
from app.services.semantic_chunker import SemanticChunker
from app.services.raptor import RaptorService
from app.services.bm25_index import get_bm25_service
from qdrant_client.models import Distance, VectorParams, PointStruct
from datetime import datetime
import hashlib
import json as json_module

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class IngestionService:
    """Orchestrates document ingestion across all knowledge layers."""

    def __init__(self):
        self.llm_service = LLMService.get_instance()

        # Embedded Qdrant - shared singleton (no server needed)
        self.qdrant_client = get_qdrant_client()
        self.collection_name = settings.QDRANT_COLLECTION

        # Thread pool for CPU-bound operations (PDF parsing, etc.)
        self.executor = ThreadPoolExecutor(max_workers=4)

        # GLiNER is loaded lazily on first ingest to keep startup fast.
        # Access via self.gliner_model property which triggers the load.
        self._gliner_model = None

        # Phase B2: Docling VLM parser (lazy-loaded, optional)
        self.docling_parser = None
        if settings.USE_DOCLING and DoclingParser.is_available():
            self.docling_parser = DoclingParser.get_instance()
            logger.info("Docling VLM parser enabled")

        # Phase B3: Semantic chunker (optional)
        self.semantic_chunker = None
        if settings.USE_SEMANTIC_CHUNKING and SemanticChunker.is_available():
            self.semantic_chunker = SemanticChunker(max_tokens=settings.SEMANTIC_CHUNK_TOKENS)
            logger.info("Semantic chunker enabled")

        # Phase B4: RAPTOR hierarchy builder (optional)
        self.raptor = None
        if settings.USE_RAPTOR and RaptorService.is_available():
            self.raptor = RaptorService(self.llm_service)
            logger.info("RAPTOR hierarchy builder enabled")

        # Atlas 3.0: BM25 sparse index
        self.bm25_service = get_bm25_service()

        # Atlas 3.0: Strict ontology for edge types
        self._allowed_edge_types = set(
            t.strip() for t in settings.GRAPH_ONTOLOGY_EDGE_TYPES.split(",") if t.strip()
        )

        logger.info("IngestionService initialized (embedded Qdrant + SQLite + BM25)")

    @property
    def gliner_model(self):
        """Lazy-load GLiNER on first access (only needed during actual ingestion)."""
        if self._gliner_model is None:
            self._gliner_model = self._load_gliner_model()
            if self._gliner_model is None:
                raise RuntimeError(
                    "GLiNER model is required for ingestion. Fix the model install/config and retry."
                )
        return self._gliner_model

    @gliner_model.setter
    def gliner_model(self, value):
        """Allow direct assignment (e.g., for testing or manual injection)."""
        self._gliner_model = value

    def _load_gliner_model(self):
        """Load GLiNER: try ONNX if available, else PyTorch."""
        from gliner import GLiNER

        gliner_path = Path(settings.MODELS_DIR) / "gliner_small-v2.1"
        onnx_path = gliner_path / "model.onnx"

        if onnx_path.exists():
            try:
                model = GLiNER.from_pretrained(str(gliner_path), load_onnx_model=True)
                logger.info(f"Loaded GLiNER ONNX model from {gliner_path}")
                return model
            except Exception as e:
                logger.warning(f"GLiNER ONNX load failed ({e}), falling back to PyTorch")

        try:
            if gliner_path.exists():
                model = GLiNER.from_pretrained(str(gliner_path))
                logger.info(f"Loaded GLiNER model from {gliner_path}")
            else:
                logger.warning(f"GLiNER model not found at {gliner_path}, downloading...")
                model = GLiNER.from_pretrained("urchade/gliner_small-v2.1")
                logger.info("Loaded GLiNER model from HuggingFace")
            return model
        except Exception as e:
            logger.error(f"GLiNER entity extraction unavailable: {e}")
            return None

    async def _ensure_collection(self):
        """Create Qdrant collection if it doesn't exist."""
        collections = self.qdrant_client.get_collections().collections
        collection_names = [c.name for c in collections]

        if self.collection_name not in collection_names:
            dimension = self.llm_service.embedding_dimension
            logger.info(f"Creating Qdrant collection with dimension {dimension}")
            self.qdrant_client.create_collection(
                collection_name=self.collection_name,
                vectors_config=VectorParams(size=dimension, distance=Distance.COSINE),
            )

    async def _embed_text(self, text: str) -> List[float]:
        """Generate embedding using bundled LLM service."""
        return await self.llm_service.embed(text)

    def _calculate_hash(self, file_path: str) -> str:
        """Calculate SHA256 hash of file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    def _get_file_type_and_mime(self, filename: str) -> Tuple[str, str]:
        """
        Determine file type and MIME type from filename.
        
        Returns:
            Tuple of (file_type, mime_type)
        """
        lower_filename = filename.lower()
        
        if lower_filename.endswith('.pdf'):
            return ('pdf', 'application/pdf')
        elif lower_filename.endswith('.docx'):
            return ('docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        elif lower_filename.endswith('.doc'):
            return ('doc', 'application/msword')
        elif lower_filename.endswith('.txt'):
            return ('txt', 'text/plain')
        else:
            # Default to text for unknown extensions
            return ('unknown', 'text/plain')

    async def ingest_document(
        self, file_path: str, filename: str, project_id: Optional[str] = None, predefined_doc_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Main ingestion pipeline.

        Args:
            file_path: Path to the document file
            filename: Original filename
            project_id: Optional project to scope this document to
            predefined_doc_id: Optional pre-generated doc_id from frontend routing

        Returns:
            Summary of ingestion with statistics
        """
        session = get_session()
        doc_id = None
        document = None
        start_time = time.time()
        try:
            logger.info(f"Starting ingestion pipeline for: {filename}")

            # Determine file type and MIME type
            file_type, mime_type = self._get_file_type_and_mime(filename)

            if predefined_doc_id:
                # Document already registered
                doc_id = predefined_doc_id
                document = session.query(Document).filter(Document.id == doc_id).first()
                if not document:
                    file_hash = self._calculate_hash(file_path)
                    file_size = Path(file_path).stat().st_size
                    document = Document(
                        id=doc_id,
                        filename=filename,
                        file_hash=file_hash,
                        file_path=file_path,
                        file_size=file_size,
                        mime_type=mime_type,
                        status="processing",
                        project_id=project_id,
                        uploaded_at=datetime.utcnow(),
                    )
                    session.add(document)
                    session.commit()
            else:
                # 1. Check for duplicate
                file_hash = self._calculate_hash(file_path)
                file_size = Path(file_path).stat().st_size

                dup_query = session.query(Document).filter(Document.file_hash == file_hash)
                if project_id:
                    dup_query = dup_query.filter(Document.project_id == project_id)
                existing_doc = dup_query.first()

                if existing_doc:
                    logger.info(f"Skipping duplicate: {filename}")
                    return {
                        "status": "duplicate",
                        "doc_id": str(existing_doc.id),
                        "message": f"Document already exists as {existing_doc.filename}",
                    }

                # Create new document
                doc_id = str(uuid.uuid4())
                document = Document(
                    id=doc_id,
                    filename=filename,
                    file_hash=file_hash,
                    file_path=file_path,
                    file_size=file_size,
                    mime_type=mime_type,
                    status="processing",
                    project_id=project_id,
                    uploaded_at=datetime.utcnow(),
                )
                session.add(document)
                session.commit()
                logger.info(f"Document created: {filename} (ID: {doc_id})")

            # 2. Extract text from document
            t0_extract = time.perf_counter()
            pages = await self._extract_text_from_file(file_path, file_type)
            logger.info(f"[INGEST TIMING] extract_text: {time.perf_counter() - t0_extract:.2f}s | Extracted {len(pages)} sections from {file_type.upper()} file")

            if len(pages) == 0:
                logger.warning(f"No text extracted from {file_type.upper()} file")
                document.status = "completed"
                document.processed_at = datetime.utcnow()
                document.total_chunks = 0
                document.processed_chunks = 0
                session.commit()
                return {
                    "status": "completed",
                    "doc_id": doc_id,
                    "filename": filename,
                    "stats": {
                        "pages": 0,
                        "chunks": 0,
                        "nodes": 0,
                        "warning": f"No text extracted from document",
                    },
                }

            # 3. Chunk document
            t0_chunk = time.perf_counter()
            chunks = self._chunk_document(pages, doc_id, filename)
            logger.info(f"[INGEST TIMING] chunk: {time.perf_counter() - t0_chunk:.2f}s | Created {len(chunks)} chunks")

            document.total_chunks = len(chunks)
            document.processed_chunks = 0
            session.commit()

            # 4. Store chunks in SQLite
            for chunk_data in chunks:
                chunk_obj = DocumentChunk(
                    id=str(uuid.uuid4()),
                    document_id=doc_id,
                    text=chunk_data["text"],
                    chunk_index=chunk_data["chunk_index"],
                    page_number=chunk_data.get("page_number"),
                    start_char=chunk_data.get("start_char"),
                    end_char=chunk_data.get("end_char"),
                    chunk_metadata=chunk_data.get("metadata", {}),
                )
                session.add(chunk_obj)
            session.commit()
            logger.info(f"Stored {len(chunks)} chunks in SQLite")

            # 5. Ensure Qdrant collection exists
            await self._ensure_collection()

            # 6. Extract entities and create graph nodes
            t0_entities = time.perf_counter()
            node_ids_by_chunk = await self._extract_entities_parallel(
                chunks, doc_id, session, document, project_id
            )
            total_nodes = sum(len(ids) for ids in node_ids_by_chunk.values())
            logger.info(f"[INGEST TIMING] entity_extraction: {time.perf_counter() - t0_entities:.2f}s | Created {total_nodes} knowledge graph nodes")
            session.commit()

            # 7. Embed and store in Qdrant
            t0_embed = time.perf_counter()
            vector_points = await self._embed_chunks_parallel(
                chunks, doc_id, node_ids_by_chunk, document, session
            )
            logger.info(f"[INGEST TIMING] embed_chunks: {time.perf_counter() - t0_embed:.2f}s | {len(vector_points)} vectors")

            # 7.5: Build RAPTOR hierarchy (Phase B4)
            raptor_points = []
            if self.raptor and vector_points and len(vector_points) >= 3:
                try:
                    embeddings = [point.vector for point in vector_points]
                    hierarchy = await self.raptor.build_hierarchy(
                        chunks=chunks,
                        embeddings=embeddings,
                        doc_id=doc_id,
                        filename=filename,
                        n_clusters=min(settings.RAPTOR_CLUSTERS, len(chunks) // 3),
                    )

                    # Store L1 cluster summaries in Qdrant
                    for summary in hierarchy.get("L1", []):
                        if not summary.get("text") or not summary.get("embedding"):
                            continue
                        raptor_points.append(PointStruct(
                            id=str(uuid.uuid4()),
                            vector=summary["embedding"],
                            payload={
                                "chunk_id": str(uuid.uuid4()),
                                "doc_id": doc_id,
                                "text": summary["text"],
                                "metadata": summary["metadata"],
                            },
                        ))

                    # Store L2 document summary in Qdrant
                    l2 = hierarchy.get("L2")
                    if l2 and l2.get("text") and l2.get("embedding"):
                        raptor_points.append(PointStruct(
                            id=str(uuid.uuid4()),
                            vector=l2["embedding"],
                            payload={
                                "chunk_id": str(uuid.uuid4()),
                                "doc_id": doc_id,
                                "text": l2["text"],
                                "metadata": l2["metadata"],
                            },
                        ))

                    logger.info(f"RAPTOR: {len(raptor_points)} hierarchy vectors generated")
                except Exception as e:
                    logger.warning(f"RAPTOR hierarchy build failed (non-fatal): {e}")

            # 8. Upload all vectors to Qdrant (L0 chunks + RAPTOR L1/L2)
            all_points = vector_points + raptor_points
            if all_points:
                t0_upsert = time.perf_counter()
                loop = asyncio.get_event_loop()
                await loop.run_in_executor(
                    self.executor,
                    lambda: self.qdrant_client.upsert(
                        collection_name=self.collection_name, points=all_points
                    ),
                )
                logger.info(f"[INGEST TIMING] qdrant_upsert: {time.perf_counter() - t0_upsert:.2f}s | Uploaded {len(all_points)} vectors to Qdrant ({len(vector_points)} chunks + {len(raptor_points)} RAPTOR)")

            # 8.5: Index chunks in BM25 for sparse retrieval (Atlas 3.0)
            try:
                self.bm25_service.add_documents(doc_id, chunks)
                logger.info(f"Added {len(chunks)} chunks to BM25 index")
            except Exception as e:
                logger.warning(f"BM25 indexing failed (non-fatal): {e}")

            # 8.6: Extract paper structure metadata (Phase 4, Task 4.1)
            try:
                full_text = "\n".join(p["text"] for p in pages)
                paper_structure = await self._extract_paper_structure(full_text, filename)
                paper_structure["page_count"] = len(pages)
                paper_structure["total_chars"] = sum(p.get("char_count", len(p["text"])) for p in pages)
                document.doc_metadata = paper_structure
                session.commit()
                logger.info(f"Extracted paper structure for {filename}: title={paper_structure.get('title', 'N/A')}")
            except Exception as e:
                logger.warning(f"Paper structure extraction failed (non-fatal): {e}")
                # Still store basic metadata even if LLM extraction fails
                document.doc_metadata = {
                    "page_count": len(pages),
                    "total_chars": sum(p.get("char_count", len(p["text"])) for p in pages),
                }
                session.commit()

            # 9. Mark as completed
            document.status = "completed"
            document.processed_at = datetime.utcnow()
            document.processed_chunks = document.total_chunks
            session.commit()

            elapsed = time.time() - start_time
            result = {
                "status": "success",
                "doc_id": doc_id,
                "filename": filename,
                "stats": {
                    "pages": len(pages),
                    "chunks": len(chunks),
                    "nodes": total_nodes,
                },
            }
            logger.info(f"[INGEST TIMING] total: {elapsed:.2f}s | Ingestion complete: {filename}")
            return result

        except Exception as e:
            logger.error(f"Ingestion failed for {filename}: {str(e)}", exc_info=True)
            try:
                session.rollback()
            except Exception:
                pass
            if document is not None:
                try:
                    document.status = "failed"
                    session.commit()
                except Exception:
                    try:
                        session.rollback()
                    except Exception:
                        pass
            return {"status": "failed", "error": str(e), "filename": filename}
        finally:
            session.close()

    # ----------------------------------------------------------------
    # Generic text extraction - dispatches based on file type
    # ----------------------------------------------------------------

    async def _extract_text_from_file(self, file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """
        Extract text from document based on file type.
        
        Args:
            file_path: Path to the document file
            file_type: Type of file ('pdf', 'docx', 'doc', 'txt', etc.)
            
        Returns:
            List of dictionaries with extracted text and metadata
        """
        if file_type == 'pdf':
            return await self._extract_pdf_text(file_path)
        elif file_type == 'docx':
            return await self._extract_docx_text(file_path)
        elif file_type in ('doc', 'unknown'):
            # For .doc files, try DOCX first (many modern .doc files are actually .docx format)
            # Then fall back to txt extraction as a last resort
            try:
                return await self._extract_docx_text(file_path)
            except Exception as e:
                logger.warning(f"DOCX extraction failed for .doc file: {e}, treating as text")
                return await self._extract_txt_text(file_path)
        elif file_type == 'txt':
            return await self._extract_txt_text(file_path)
        else:
            logger.warning(f"Unknown file type: {file_type}, attempting text extraction")
            return await self._extract_txt_text(file_path)

    # ----------------------------------------------------------------
    # PDF extraction
    # ----------------------------------------------------------------

    async def _extract_pdf_text(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from PDF with page numbers (async).

        Phase B2: Tries Docling VLM parser first for structure-preserving
        extraction (tables, charts), falls back to pdfplumber/PyPDF.
        """
        # Phase B2: Try Docling first (preserves tables and structure)
        if self.docling_parser:
            try:
                loop = asyncio.get_event_loop()
                pages = await loop.run_in_executor(
                    self.executor,
                    self.docling_parser.parse_document,
                    file_path,
                )
                if pages:
                    logger.info(f"Docling extracted {len(pages)} pages")
                    return pages
            except Exception as e:
                logger.warning(f"Docling failed, falling back to pdfplumber: {e}")

        # Fallback: pdfplumber / PyPDF
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, self._extract_pdf_text_sync, file_path)

    def _extract_pdf_text_sync(self, file_path: str) -> List[Dict[str, Any]]:
        """Synchronous PDF text extraction."""
        pages = []

        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, start=1):
                        text = page.extract_text()
                        if text and text.strip():
                            pages.append({"page_number": page_num, "text": text, "char_count": len(text)})
                if pages:
                    return pages
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}, trying PyPDF...")

        try:
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text and text.strip():
                    pages.append({"page_number": page_num, "text": text, "char_count": len(text)})
        except Exception as e:
            logger.error(f"PyPDF extraction also failed: {e}")

        return pages

    # ----------------------------------------------------------------
    # DOCX extraction
    # ----------------------------------------------------------------

    async def _extract_docx_text(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from DOCX file (async wrapper)."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, self._extract_docx_text_sync, file_path)

    def _extract_docx_text_sync(self, file_path: str) -> List[Dict[str, Any]]:
        """Synchronous DOCX text extraction."""
        pages = []
        
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available. Install it to process DOCX files.")
            return pages
        
        try:
            doc = DocxDocument(file_path)
            full_text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text += cell.text + "\n"
            
            if full_text.strip():
                # Treat entire DOCX as one page for consistency
                pages.append({"page_number": 1, "text": full_text.strip(), "char_count": len(full_text)})
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
        
        return pages

    # ----------------------------------------------------------------
    # TXT extraction
    # ----------------------------------------------------------------

    async def _extract_txt_text(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from TXT file (async wrapper)."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(self.executor, self._extract_txt_text_sync, file_path)

    def _extract_txt_text_sync(self, file_path: str) -> List[Dict[str, Any]]:
        """Synchronous TXT text extraction."""
        pages = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            if text.strip():
                # Treat entire text file as one page
                pages.append({"page_number": 1, "text": text.strip(), "char_count": len(text)})
        except UnicodeDecodeError:
            # Try with a different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                
                if text.strip():
                    pages.append({"page_number": 1, "text": text.strip(), "char_count": len(text)})
            except Exception as e:
                logger.error(f"TXT extraction with fallback encoding failed: {e}")
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
        
        return pages

    # ----------------------------------------------------------------
    # Chunking
    # ----------------------------------------------------------------

    def _chunk_document(
        self, pages: List[Dict[str, Any]], doc_id: str, filename: str
    ) -> List[Dict[str, Any]]:
        """Chunk document using semantic or fixed-size strategy.

        Phase B3: Uses semantic chunker if available, otherwise falls back
        to the original fixed-size overlap chunking.
        """
        if self.semantic_chunker:
            try:
                chunks = self.semantic_chunker.chunk_pages(pages, doc_id, filename)
                if chunks:
                    logger.info(f"Semantic chunking produced {len(chunks)} chunks")
                    return chunks
            except Exception as e:
                logger.warning(f"Semantic chunking failed, falling back to fixed-size: {e}")

        return self._chunk_document_fixed_size(pages, doc_id, filename)

    def _chunk_document_fixed_size(
        self, pages: List[Dict[str, Any]], doc_id: str, filename: str
    ) -> List[Dict[str, Any]]:
        """Original fixed-size chunking with overlap (fallback)."""
        chunks = []
        chunk_index = 0

        for page in pages:
            text = page["text"]
            page_num = page["page_number"]
            start = 0
            while start < len(text):
                end = start + settings.CHUNK_SIZE
                chunk_text = text[start:end]

                if end < len(text):
                    last_period = chunk_text.rfind(".")
                    last_newline = chunk_text.rfind("\n")
                    break_point = max(last_period, last_newline)
                    if break_point > settings.CHUNK_SIZE * 0.5:
                        end = start + break_point + 1
                        chunk_text = text[start:end]

                chunks.append(
                    {
                        "text": chunk_text.strip(),
                        "chunk_index": chunk_index,
                        "page_number": page_num,
                        "start_char": start,
                        "end_char": end,
                        "metadata": {"filename": filename, "doc_id": doc_id, "page": page_num},
                    }
                )
                chunk_index += 1
                start = end - settings.CHUNK_OVERLAP

        return chunks

    # ----------------------------------------------------------------
    # Entity extraction
    # ----------------------------------------------------------------

    async def _extract_entities_parallel(
        self,
        chunks: List[Dict[str, Any]],
        doc_id: str,
        session: Session,
        document: Document,
        project_id: Optional[str] = None,
    ) -> Dict[int, List[str]]:
        """Extract entities from all chunks in parallel."""
        node_ids_by_chunk: Dict[int, List[str]] = {}
        entity_labels = [
            l.strip() for l in settings.GRAPH_ENTITY_LABELS.split(",") if l.strip()
        ]

        extraction_tasks = []
        for chunk in chunks:
            task = self._extract_entities_gliner(chunk["text"], labels=entity_labels)
            extraction_tasks.append((chunk, task))

        results = await asyncio.gather(
            *[task for _, task in extraction_tasks], return_exceptions=True
        )

        processed_count = 0
        for (chunk, _), entities_or_error in zip(extraction_tasks, results):
            if isinstance(entities_or_error, Exception):
                logger.error(f"Entity extraction failed for chunk {chunk['chunk_index']}: {entities_or_error}")
                raise entities_or_error

            entities = entities_or_error
            chunk_uuid = str(uuid.uuid5(uuid.UUID(doc_id), f"chunk-{chunk['chunk_index']}"))

            chunk_node_ids: List[str] = []
            entity_nodes = []

            for entity in entities:
                node_id = str(uuid.uuid4())
                node = Node(
                    id=node_id,
                    label=entity.get("type", "Entity"),
                    document_id=doc_id,
                    project_id=project_id,
                    properties={
                        "name": entity["name"],
                        "description": entity.get("description", ""),
                        "chunk_id": chunk_uuid,
                        "page_number": chunk["page_number"],
                        "confidence": entity.get("confidence", 0.8),
                    },
                )
                session.add(node)
                chunk_node_ids.append(node_id)
                entity_nodes.append(node)

            if chunk_node_ids:
                node_ids_by_chunk[chunk["chunk_index"]] = chunk_node_ids

            # Atlas 3.0: Evidence-bound relationship extraction
            # Instead of naive CO_OCCURS edges, use LLM to extract typed relationships
            # with evidence quotes that are validated by a critic step
            if len(entity_nodes) > 1 and settings.ENABLE_EVIDENCE_BOUND_EXTRACTION:
                try:
                    entity_names = [
                        (n.id, (n.properties or {}).get("name", "Unknown"))
                        for n in entity_nodes
                    ]
                    relationships = await self._extract_evidence_bound_relationships(
                        chunk["text"], entity_names, chunk_uuid
                    )
                    for rel in relationships:
                        edge = Edge(
                            id=str(uuid.uuid4()),
                            source_id=rel["source_id"],
                            target_id=rel["target_id"],
                            type=rel["type"],
                            document_id=doc_id,
                            project_id=project_id,
                            properties={
                                "chunk_id": chunk_uuid,
                                "evidence_quote": rel.get("evidence_quote", ""),
                                "confidence": rel.get("confidence", 0.5),
                                "context": chunk["text"][:200],
                            },
                        )
                        session.add(edge)
                except Exception as e:
                    logger.warning(f"Evidence-bound extraction failed for chunk {chunk['chunk_index']}, "
                                   f"falling back to CO_OCCURS: {e}")
                    # Fallback: create CO_OCCURS edges
                    for i, node1 in enumerate(entity_nodes):
                        for node2 in entity_nodes[i + 1:]:
                            edge = Edge(
                                id=str(uuid.uuid4()),
                                source_id=node1.id,
                                target_id=node2.id,
                                type="CO_OCCURS",
                                document_id=doc_id,
                                project_id=project_id,
                                properties={"chunk_id": chunk_uuid, "context": chunk["text"][:200]},
                            )
                            session.add(edge)
            elif len(entity_nodes) > 1:
                # Legacy mode: CO_OCCURS edges when evidence-bound extraction is disabled
                for i, node1 in enumerate(entity_nodes):
                    for node2 in entity_nodes[i + 1:]:
                        edge = Edge(
                            id=str(uuid.uuid4()),
                            source_id=node1.id,
                            target_id=node2.id,
                            type="CO_OCCURS",
                            document_id=doc_id,
                            project_id=project_id,
                            properties={"chunk_id": chunk_uuid, "context": chunk["text"][:200]},
                        )
                        session.add(edge)

            processed_count += 1
            # Update progress more frequently for better UX
            # Calculate optimal update frequency based on total chunks (at least every 3 chunks or 10%)
            update_frequency = max(3, len(chunks) // 10) if chunks else 1
            if processed_count % update_frequency == 0 or processed_count == len(chunks):
                try:
                    session.refresh(document)
                    document.processed_chunks = processed_count
                    session.commit()
                    if processed_count % (update_frequency * 2) == 0 or processed_count == len(chunks):
                        logger.info(f"Entity extraction progress: {processed_count}/{len(chunks)} chunks")
                except Exception as e:
                    logger.warning(f"Failed to update progress: {e}")
                    session.rollback()

        return node_ids_by_chunk

    async def _embed_chunks_parallel(
        self,
        chunks: List[Dict[str, Any]],
        doc_id: str,
        node_ids_by_chunk: Dict[int, List[str]],
        document: Document,
        session: Session,
    ) -> List[PointStruct]:
        """Embed all chunks using batch processing (faster & thread-safe)."""
        if not chunks:
            return []
            
        texts = [chunk["text"] for chunk in chunks]
        
        try:
            # Use batch embedding instead of individual calls
            embeddings = await self.llm_service.embed_batch(texts)
        except Exception as e:
            logger.error(f"Batch embedding failed: {e}")
            raise

        vector_points = []
        
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_uuid = str(uuid.uuid5(uuid.UUID(doc_id), f"chunk-{chunk['chunk_index']}"))

            metadata = chunk.get("metadata", {})
            if chunk["chunk_index"] in node_ids_by_chunk:
                metadata["node_ids"] = node_ids_by_chunk[chunk["chunk_index"]]

            point = PointStruct(
                id=chunk_uuid,
                vector=embedding,
                payload={
                    "chunk_id": chunk_uuid,
                    "doc_id": doc_id,
                    "text": chunk["text"],
                    "metadata": metadata,
                },
            )
            vector_points.append(point)

        # Update progress at the end (since batch is all-or-nothing efficiently)
        try:
            session.refresh(document)
            document.processed_chunks = len(chunks)
            session.commit()
            logger.info(f"Embedding progress: {len(chunks)}/{len(chunks)} chunks (batch)")
        except Exception as e:
            logger.warning(f"Failed to update progress: {e}")
            session.rollback()

        return vector_points

    async def _extract_entities_gliner(
        self, text: str, labels: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """Extract entities using GLiNER model."""
        if labels is None:
            labels = [
                l.strip() for l in settings.GRAPH_ENTITY_LABELS.split(",") if l.strip()
            ]

        if self.gliner_model is None:
            raise RuntimeError("GLiNER model not available")

        try:
            loop = asyncio.get_event_loop()
            window_size = 1200
            overlap = 150
            windows = []
            if len(text) <= window_size:
                windows = [text]
            else:
                start = 0
                while start < len(text):
                    end = min(len(text), start + window_size)
                    windows.append(text[start:end])
                    if end == len(text):
                        break
                    start = end - overlap

            def _predict_all():
                all_preds = []
                for window_text in windows:
                    all_preds.extend(
                        self.gliner_model.predict_entities(window_text, labels=labels, threshold=0.2)
                    )
                return all_preds

            predictions = await loop.run_in_executor(self.executor, _predict_all)

            entities = []
            seen_entities = set()
            type_mapping = {
                "Person": "person", "Organization": "organization", "Location": "location",
                "Concept": "concept", "Method": "method", "Chemical": "chemical",
                "Date": "date", "Event": "event", "Work": "work",
                "Title": "title", "Institution": "institution",
            }

            for pred in predictions:
                entity_name = pred.get("text", "").strip()
                entity_type = pred.get("label", "Concept")
                confidence = pred.get("score", 0.5)

                if not entity_name or len(entity_name) < 2:
                    continue

                entity_key = (entity_name.lower(), entity_type)
                if entity_key in seen_entities:
                    continue
                seen_entities.add(entity_key)

                entities.append(
                    {
                        "name": entity_name,
                        "type": type_mapping.get(entity_type, "concept"),
                        "description": "",
                        "confidence": float(confidence),
                    }
                )

            return entities

        except Exception as e:
            logger.error(f"GLiNER extraction failed: {str(e)}", exc_info=True)
            raise RuntimeError(f"GLiNER entity extraction failed: {str(e)}") from e

    # ----------------------------------------------------------------
    # Atlas 3.0: Evidence-Bound Relationship Extraction (Phase 2)
    # ----------------------------------------------------------------

    async def _extract_evidence_bound_relationships(
        self,
        chunk_text: str,
        entity_names: List[tuple],  # [(node_id, entity_name), ...]
        chunk_id: str,
    ) -> List[Dict[str, Any]]:
        """Extract typed, evidence-bound relationships between entities using LLM.

        Instead of naive co-occurrence edges, this method:
        1. Prompts the LLM to identify relationships with evidence quotes
        2. Validates that evidence quotes are actual substrings of the chunk
        3. Filters to only allowed ontology edge types

        Args:
            chunk_text: The text of the chunk containing the entities.
            entity_names: List of (node_id, entity_name) tuples.
            chunk_id: UUID of the chunk.

        Returns:
            List of validated relationship dicts.
        """
        if len(entity_names) < 2:
            return []

        # Build entity list for prompt
        entity_list = ", ".join(name for _, name in entity_names)
        allowed_types = ", ".join(sorted(self._allowed_edge_types))

        prompt = f"""Analyze this text and identify relationships between the given entities.

ENTITIES: {entity_list}

ALLOWED RELATIONSHIP TYPES: {allowed_types}

TEXT:
{chunk_text[:2000]}

Return ONLY a JSON array of relationships. Each relationship MUST include:
- "source": exact entity name from the list
- "target": exact entity name from the list
- "type": one of the allowed relationship types above
- "evidence_quote": the EXACT substring from the text that proves this relationship

Rules:
- Only include relationships explicitly stated in the text
- The evidence_quote MUST be a direct quote from the text (copy-paste)
- Do NOT infer relationships that are not stated
- Return an empty array [] if no clear relationships exist

JSON array:"""

        try:
            response = await self.llm_service.generate(
                prompt=prompt, temperature=0.1, max_tokens=1024
            )

            # Parse JSON response
            relationships_raw = self._parse_relationships_json(response)
        except Exception as e:
            logger.debug(f"Relationship extraction LLM call failed: {e}")
            return []

        # Build name -> node_id map
        name_to_id = {name.lower(): nid for nid, name in entity_names}

        validated = []
        for rel in relationships_raw:
            source_name = rel.get("source", "").lower()
            target_name = rel.get("target", "").lower()
            rel_type = rel.get("type", "").upper()
            evidence = rel.get("evidence_quote", "")

            # Validate source and target exist
            source_id = name_to_id.get(source_name)
            target_id = name_to_id.get(target_name)
            if not source_id or not target_id or source_id == target_id:
                continue

            # Validate edge type is in ontology
            if rel_type not in self._allowed_edge_types:
                # Try to map to closest allowed type
                rel_type = "RELATED_TO"

            # Atlas 3.0 Critic: Validate evidence quote is a real substring
            if settings.ENABLE_GRAPH_CRITIC and evidence:
                if not self._validate_evidence_quote(chunk_text, evidence):
                    logger.debug(f"Dropping edge {source_name}->{target_name}: evidence not found in text")
                    continue

            validated.append({
                "source_id": source_id,
                "target_id": target_id,
                "type": rel_type,
                "evidence_quote": evidence,
                "confidence": 0.8 if evidence else 0.5,
            })

        return validated

    @staticmethod
    def _validate_evidence_quote(chunk_text: str, evidence_quote: str) -> bool:
        """Critic validation: Check if the evidence quote exists in the chunk text.

        This is the core anti-hallucination mechanism. If the LLM fabricated
        the evidence quote, the edge is dropped.

        Args:
            chunk_text: The original chunk text.
            evidence_quote: The claimed evidence quote from the LLM.

        Returns:
            True if the quote (or a close substring) exists in the text.
        """
        if not evidence_quote or len(evidence_quote) < 10:
            return False

        # Normalize whitespace for comparison
        normalized_text = " ".join(chunk_text.lower().split())
        normalized_quote = " ".join(evidence_quote.lower().split())

        # Exact substring match
        if normalized_quote in normalized_text:
            return True

        # Fuzzy match: check if at least 80% of the quote words appear in sequence
        quote_words = normalized_quote.split()
        if len(quote_words) < 3:
            return normalized_quote in normalized_text

        # Sliding window match
        text_words = normalized_text.split()
        window_size = len(quote_words)
        match_threshold = int(window_size * 0.7)

        for i in range(len(text_words) - window_size + 1):
            window = text_words[i:i + window_size]
            matches = sum(1 for qw, tw in zip(quote_words, window) if qw == tw)
            if matches >= match_threshold:
                return True

        return False

    @staticmethod
    def _parse_relationships_json(response: str) -> List[Dict[str, Any]]:
        """Parse JSON array of relationships from LLM response."""
        text = response.strip()

        # Strip markdown code blocks
        if text.startswith("```"):
            import re
            text = re.sub(r"^```(?:json)?\s*\n?", "", text)
            text = re.sub(r"\n?```\s*$", "", text)

        # Try direct parse
        try:
            result = json_module.loads(text)
            if isinstance(result, list):
                return result
            if isinstance(result, dict) and "relationships" in result:
                return result["relationships"]
            return []
        except json_module.JSONDecodeError:
            pass

        # Try to find JSON array in response
        import re
        match = re.search(r'\[[\s\S]*\]', text)
        if match:
            try:
                return json_module.loads(match.group())
            except json_module.JSONDecodeError:
                pass

        return []

    # ----------------------------------------------------------------
    # Paper Structure Extraction (Phase 4, Task 4.1)
    # ----------------------------------------------------------------

    async def _extract_paper_structure(self, text: str, filename: str) -> Dict[str, Any]:
        """Extract structured academic paper metadata using LLM.

        Analyzes the first ~3000 chars of a document to extract title, authors,
        year, abstract, methodology, key findings, and paper type.

        Returns:
            Dict with extracted metadata fields. Missing fields default to empty.
        """
        import re
        import json as json_module

        # Use first 3000 chars for structure extraction (covers most abstracts/intros)
        sample = text[:3000]

        prompt = f"""Extract the following from this academic paper/document. Return ONLY valid JSON, no other text.
{{
    "title": "paper title or document title",
    "authors": ["author 1", "author 2"],
    "year": 2024,
    "abstract": "abstract text (first 500 chars max)",
    "methodology": "brief description of methods used",
    "key_findings": ["finding 1", "finding 2"],
    "limitations": ["limitation 1"],
    "paper_type": "empirical|review|theoretical|meta-analysis|report|other"
}}

If any field cannot be determined, use empty string or empty list.

Document text (first 3000 chars):
{sample}

JSON:"""

        response = await self.llm_service.generate(
            prompt=prompt, temperature=0.1, max_tokens=1024
        )

        # Parse JSON from LLM response
        result = self._parse_structure_response(response)

        # Fallback: extract title from filename if LLM didn't get it
        if not result.get("title"):
            # Strip extension and clean up filename
            name = Path(filename).stem
            result["title"] = name.replace("_", " ").replace("-", " ").strip()

        return result

    @staticmethod
    def _parse_structure_response(response: str) -> Dict[str, Any]:
        """Parse JSON from LLM response, handling code blocks and malformed output."""
        import re
        import json as json_module

        # Try direct JSON parse
        text = response.strip()

        # Strip markdown code blocks if present
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*\n?", "", text)
            text = re.sub(r"\n?```\s*$", "", text)

        try:
            return json_module.loads(text)
        except json_module.JSONDecodeError:
            pass

        # Try to find JSON object in the response
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            try:
                return json_module.loads(match.group())
            except json_module.JSONDecodeError:
                pass

        # Return empty structure on complete failure
        return {
            "title": "",
            "authors": [],
            "year": None,
            "abstract": "",
            "methodology": "",
            "key_findings": [],
            "limitations": [],
            "paper_type": "other",
        }
